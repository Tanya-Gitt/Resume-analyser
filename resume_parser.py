import io
from pdfminer.high_level import extract_text as _pdf_extract
from docx import Document


def extract_text_from_file(file) -> str | None:
    """Extract plain text from a PDF, DOCX, or TXT upload."""
    try:
        file_type = file.type
        raw = file.read()

        if file_type == "application/pdf":
            text = _pdf_extract(io.BytesIO(raw))
        elif file_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ):
            doc = Document(io.BytesIO(raw))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            # Also grab text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text)
            text = "\n".join(paragraphs)
        elif file_type == "text/plain":
            text = raw.decode("utf-8", errors="replace")
        else:
            return None

        text = text.strip()
        return text if text else None

    except Exception as e:
        print(f"[resume_parser] Error reading {getattr(file, 'name', '?')}: {e}")
        return None
